import os
import openpyxl
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import win32com.client
import tkinter.font as tkFont


class FolderPriorityReader:
    def __init__(self, excel_path):
        self.excel_path = excel_path
        self.level_dicts = {}     # 各级的名字 -> 优先级
        self.level_names = []     

    def read_priority(self):
        import openpyxl
        wb = openpyxl.load_workbook(self.excel_path)
        sheet = wb.active

        self.level_dicts.clear()
        self.level_names.clear()

        for col in range(1, sheet.max_column + 1, 2):
            header = sheet.cell(row=1, column=col).value
            if not header or "级文件夹" not in str(header):
                continue
            self.level_names.append(header.strip())
            level_dict = {}
            for row in range(3, sheet.max_row + 1):
                name = sheet.cell(row=row, column=col).value
                prio = sheet.cell(row=row, column=col + 1).value
                if name and prio:
                    try:
                        # level_dict[str(name).strip()] = int(prio)
                        level_dict[str(name).strip().upper()] = int(prio) # 统一将小写转换为大写
                    except ValueError:
                        continue
            self.level_dicts[header.strip()] = level_dict
        return self.level_dicts, self.level_names



def get_sort_key_from_path(path: str, priority_dicts: dict, level_names: list, root: str) -> tuple:
    import os
    path = os.path.abspath(path)
    root = os.path.abspath(root)
    relative_path = os.path.relpath(path, root)
    parts = relative_path.split(os.sep)

    # 拿出路径中与级别匹配的部分 + 文件名（不含扩展名）作为最后一级
    folder_parts = parts[:-1][-len(level_names):]
    filename_part = os.path.splitext(os.path.basename(path))[0]
    folder_parts.append(filename_part)

    sort_key = []
    for i, level in enumerate(level_names):
        name = folder_parts[i] if i < len(folder_parts) else ''
        # prio = priority_dicts.get(level, {}).get(name, 999)
        prio = priority_dicts.get(level, {}).get(name.upper(), 999) # 小写转换为大写
        sort_key.append(prio)
    return tuple(sort_key)



class WordImageInserter:
    # def __init__(self, image_root: str, output_path: str, excel_path: str, template_path: str):
    def __init__(self, image_root: str, output_path: str, excel_path: str, template_path: str, log_func=None):

        self.image_root = image_root
        self.output_path = output_path
        self.template_path = template_path
        self.doc = Document(template_path)
        # self.folder_priorities = FolderPriorityReader(excel_path).read_priority()

        self.log_func = log_func
        reader = FolderPriorityReader(excel_path)
        self.folder_priorities, self.level_names = reader.read_priority()
        


    def convert_path(self, path: str) -> str:
        return path.replace("\\", "_").replace("/", "_")

    '''
    def get_all_images(self):
        valid_exts = [".jpg", ".jpeg", ".png", ".bmp"]
        image_paths = []
        for root, _, files in os.walk(self.image_root):
            for f in files:
                if os.path.splitext(f)[1].lower() in valid_exts:
                    image_paths.append(os.path.join(root, f))
        # return sorted(image_paths, key=lambda p: get_sort_key_from_path(p, self.folder_priorities))
        # return sorted(image_paths, key=lambda p: get_sort_key_from_path(p, self.folder_priorities, self.image_root))
        # 添加调试打印
        for img_path in image_paths:
            # sort_key = get_sort_key_from_path(img_path, self.folder_priorities, self.image_root)
            sort_key = get_sort_key_from_path(img_path, self.folder_priorities, self.level_names, self.image_root)
            print(f"图片：{img_path}")
            print(f"  -> 排序 key：{sort_key}")

        # return sorted(image_paths, key=lambda p: get_sort_key_from_path(p, self.folder_priorities, self.image_root))
        return sorted(
            image_paths,
            key=lambda p: get_sort_key_from_path(p, self.folder_priorities, self.level_names, self.image_root)
        )'''

    def get_all_images(self):
        valid_exts = [".jpg", ".jpeg", ".png", ".bmp"]
        image_paths = []
        for root, _, files in os.walk(self.image_root):
            for f in files:
                if os.path.splitext(f)[1].lower() in valid_exts:
                    image_paths.append(os.path.join(root, f))

        image_with_keys = []
        for img_path in image_paths:
            sort_key = get_sort_key_from_path(img_path, self.folder_priorities, self.level_names, self.image_root)
            image_with_keys.append((img_path, sort_key))

            # 始终输出到终端
            print(f"{img_path} \n-> 排序 key: {sort_key}")

            # 如果排序 key 中含有 999，说明可能有拼写错误，需要提示用户
            if 999 in sort_key:
                warning = (
                    f"*******************************************\n"
                    f"（注意！）{img_path} -> 排序 key: {sort_key}\n"
                    f"该路径中可能存在文件名拼写错误或未在‘文件名_级别.xlsx’中声明该文件，请检查。\n"
                    f"*******************************************\n"
                )
                if self.log_func:
                    self.log_func(warning)
                else:
                    print(warning)


        return [img for img, _ in sorted(image_with_keys, key=lambda x: x[1])]



        '''
    def read_image_name_map(self):
        wb = openpyxl.load_workbook("图片命名映射表.xlsx")
        sheet = wb.active
        name_map = {}
        for row in sheet.iter_rows(min_row=2, max_col=2):
            original, mapped = row[0].value, row[1].value
            if original and mapped:
                name_map[str(original).strip()] = str(mapped).strip()
        return name_map
    '''

    def read_image_name_map(self): 
        # 添加将小写字母统一转换成大写的功能
        wb = openpyxl.load_workbook("图片命名映射表.xlsx")
        sheet = wb.active
        name_map = {}
        for row in sheet.iter_rows(min_row=2, max_col=2):
            original, mapped = row[0].value, row[1].value
            if original and mapped:
                name_map[str(original).strip().upper()] = str(mapped).strip()
        return name_map


       
    def update_fields_with_word(self):
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False  # 静默操作
            word.Options.UpdateFieldsAtPrint = True  # 确保打印时更新字段
            doc = word.Documents.Open(os.path.abspath(self.output_path))
        
            # 更新字段并重新分页
            doc.Fields.Update()
            doc.Repaginate()  # 关键：重新计算分页
            doc.Fields.Update()
        
            # 强制整个文档刷新布局
            doc.Range().ComputeStatistics(2)  # wdStatisticPages
        
            doc.Save()
            doc.Close()
            word.Quit()
        except Exception as e:
            print(f"无法自动更新 Word 字段：{e}")
            if 'doc' in locals():
                doc.Close()
            word.Quit()

    def add_field_code(self, paragraph, field_code, font_name="Times New Roman", font_size=14):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        run = paragraph.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)
        r_element = run._r

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_code

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        # 创建字段结果 + 设置其样式
        fldResult_r = OxmlElement('w:r')
        fldResult_rPr = OxmlElement('w:rPr')

        # 字体大小（单位为1/2pt）
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        fldResult_rPr.append(sz)

        # 字体名称
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        fldResult_rPr.append(rFonts)

        fldResult_text = OxmlElement('w:t')
        fldResult_text.text = "1"  # 占位内容

        fldResult_r.append(fldResult_rPr)
        fldResult_r.append(fldResult_text)

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldResult_r)
        r_element.append(fldChar3)


    def insert_images_in_grid(self, images: list, images_per_page: int = 6):
        for idx, img_path in enumerate(images):
            if idx % images_per_page == 0:
                if idx != 0:
                    self.doc.add_page_break()

                # 插入标题
                title1 = self.doc.add_paragraph()
                run1 = title1.add_run("Microwave Radio Unit")
                run1.font.name = "Times New Roman"
                run1.font.size = Pt(18)
                run1.bold = True
                title1.alignment = WD_ALIGN_PARAGRAPH.CENTER

                title2 = self.doc.add_paragraph()
                run2 = title2.add_run("Annex B: Test screenshot")
                run2.font.name = "Times New Roman"
                run2.font.size = Pt(18)
                run2.bold = True
                title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

                title3 = self.doc.add_paragraph()
                title3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run_static = title3.add_run("Report number:\t24BXXXXXXX\t\tTotal pages: ")
                run_static.font.name = "Times New Roman"
                run_static.font.size = Pt(14)

                self.add_field_code(title3, "NUMPAGES", font_name="Times New Roman", font_size=14)  # 动态总页数
                run_static2 = title3.add_run("    Page ")
                run_static2.font.name = "Times New Roman"
                run_static2.font.size = Pt(14)

                self.add_field_code(title3, "PAGE", font_name="Times New Roman", font_size=14)  # 动态当前页

                # 表格开始
                table = self.doc.add_table(rows=3, cols=2)
                table.autofit = False
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for row in table.rows:
                    for cell in row.cells:
                        cell.width = Inches(3.45)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            row = (idx % images_per_page) // 2
            col = (idx % images_per_page) % 2
            cell = table.cell(row, col)
            cell.text = ""

            pic_p = cell.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.paragraph_format.space_before = Pt(2)
            pic_p.paragraph_format.space_after = Pt(2)
            pic_p.add_run().add_picture(img_path, width=Inches(3.425))

            relative_path = os.path.relpath(img_path, self.image_root)
            parts = os.path.splitext(relative_path)[0].split(os.sep)[-len(self.level_names):]  # 动态路径长度
            label = "_".join(parts)

            # 提取图片原始文件名（无扩展名），并查映射
            filename = parts[-1]
            # mapped_name = self.image_name_map.get(filename, filename)  # 如果找不到映射就保留原名
            mapped_name = self.image_name_map.get(filename.upper(), filename) # 增加小写字母自动转换成大写功能


            # 替换 label 中的原名为映射名
            if filename in label:
                label = label.replace(filename, mapped_name)

            fig_number = idx + 1
            desc_p = cell.add_paragraph()
            desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_run = desc_p.add_run(f"Fig. {fig_number} {label}")
            desc_run.font.name = "Times New Roman"
            desc_run.font.size = Pt(10.5)


    def generate_word(self):

        self.image_name_map = self.read_image_name_map()

        print("--- Excel 优先级字典 ---")
        for level, mapping in self.folder_priorities.items():
            print(f"{level}: {mapping}")


        images = self.get_all_images()
        self.insert_images_in_grid(images)
        self.doc.settings.element.set(qn('w:updateFields'), 'true')
        self.doc.save(self.output_path)
        self.update_fields_with_word()
        return self.output_path


class WordImageGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("图片写入文档程序")
        # self.root.geometry()

        self.default_font = tkFont.nametofont("TkDefaultFont")
        self.default_font.configure(size=14)
        self.text_font = tkFont.Font(family="TkFixedFont", size=14)

        # 设置 Treeview 样式字体
        style = ttk.Style()
        style.configure("Treeview", font=self.default_font)

        self.excel_path = os.path.join(os.path.dirname(__file__), "文件名_级别.xlsx")
        # self.template_path = os.path.join(os.path.dirname(__file__), "Annex B-8G.docx")
        self.template_path = os.path.join(os.path.dirname(__file__), "template.docx")
        # self.image_root = os.path.join(os.path.dirname(__file__), "data")
        self.image_root = os.path.join(os.path.dirname(__file__), "data", "SRU2 38G-20240801")
        self.output_path = os.path.join(os.path.dirname(__file__), "output.docx")

        self.reader = FolderPriorityReader(self.excel_path)
        self.priority_data = {}
        self.level_names = []

        self.tree = ttk.Treeview(root, columns=("Level", "Name", "Priority"), show="headings")
        # self.tree.configure(font=self.default_font)
        self.tree.heading("Level", text="等级")
        self.tree.heading("Name", text="文件夹名")
        self.tree.heading("Priority", text="优先级")
        self.tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        button_frame = tk.Frame(root)
        button_frame.pack(pady=5)
        # tk.Button(button_frame, text="读取优先级", command=self.load_excel).pack(side=tk.LEFT, padx=10)
        tk.Button(button_frame, text="重新读取优先级", font=self.default_font, command=self.load_excel).pack(side=tk.LEFT, padx=10)
        tk.Button(button_frame, text="   写入文档   ", font=self.default_font, command=self.write_docx).pack(side=tk.LEFT, padx=10)
        # tk.Button(button_frame, text="退出", command=root.quit).pack(side=tk.LEFT, padx=10)

        self.log = tk.Text(root, height=8, font=self.text_font)
        self.log.pack(fill=tk.BOTH, padx=10, pady=5)

        self.load_excel()


    def log_msg(self, text):
        self.log.insert(tk.END, text + "\n")
        self.log.see(tk.END)

    def load_excel(self):
        try:
            self.tree.delete(*self.tree.get_children())
            self.priority_data, self.level_names = self.reader.read_priority()  # 拆出级数信息
            for level, items in self.priority_data.items():
                for name, prio in items.items():
                    self.tree.insert("", tk.END, values=(level, name, prio))

            # 加入提示：显示读取了多少级
            self.log_msg(f"成功读取优先级数据，共读取到 {len(self.level_names)} 级文件夹")
        except Exception as e:
            self.log_msg(f"错误：{e}")
            messagebox.showerror("错误", f"读取 Excel 时出错：{e}")


            '''
    def write_docx(self):
        try:
            inserter = WordImageInserter(
                image_root=self.image_root,
                output_path=self.output_path,
                excel_path=self.excel_path,
                template_path=self.template_path
            )
            image_name_map = inserter.read_image_name_map()
            path = inserter.generate_word()
            # self.log_msg(f"Word 写入完成：{path}")

            for original, mapped in image_name_map.items():
                self.log_msg(f"{original} -> {mapped}")

                # 传给 WordImageInserter 实例（确保写入时能用）
                inserter.image_name_map = image_name_map

                path = inserter.generate_word()
                self.log_msg(f"Word 写入完成：{path}")

        except Exception as e:
            self.log_msg(f"错误：{e}")
            messagebox.showerror("错误", f"生成 Word 文档时出错：{e}")
            '''
    def write_docx(self):
        self.log_msg("正在写入文档，请稍等……")
        self.root.update_idletasks()  
        try:
            inserter = WordImageInserter(
                image_root=self.image_root,
                output_path=self.output_path,
                excel_path=self.excel_path,
                template_path=self.template_path,
                log_func=self.log_msg 
            )
            # 读取映射表
            image_name_map = inserter.read_image_name_map()

            # 在 GUI 日志框显示映射内容
            self.log_msg("成功读取图片命名映射表：")
            for original, mapped in image_name_map.items():
                self.log_msg(f"{original} -> {mapped}")

            # 传给 WordImageInserter 实例（确保写入时能用）
            inserter.image_name_map = image_name_map

            path = inserter.generate_word()
            self.log_msg(f"Word 写入完成：{path}")
        except Exception as e:
            self.log_msg(f"错误：{e}")
            messagebox.showerror("错误", f"生成 Word 文档时出错：{e}")



if __name__ == "__main__":
    root = tk.Tk()
    app = WordImageGUI(root)
    root.mainloop()